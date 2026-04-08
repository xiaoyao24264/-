# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), kwargs[edge].get('sz', '4'))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def add_page_header(doc, text):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_page_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('第')
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run2 = p.add_run('页 共')
    run2.font.name = '宋体'
    run2.font.size = Pt(9)
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)
    run2._r.append(instrText2)
    run2._r.append(fldChar4)
    run3 = p.add_run('页')
    run3.font.name = '宋体'
    run3.font.size = Pt(9)
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_run_font(run, font_name='宋体', font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold

doc = Document()

# 页边距
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

add_page_header(doc, '新建精密自动化设备产业园设计采购施工(EPC)总承包项目')
add_page_footer(doc)

# 标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('新建精密自动化产业园项目一季度经济活动分析报告')
set_run_font(run, '黑体', 22, True)

doc.add_paragraph()

# 一、工程基本情况
p = doc.add_paragraph()
run = p.add_run('一、工程基本情况')
set_run_font(run, '黑体', 14, True)

basic_info = [
    ('项目名称', '新建精密自动化设备产业园设计采购施工(EPC)总承包项目'),
    ('建筑面积', '71610.94 m²'),
    ('合同开工日期', '2024/10/20'),
    ('合同竣工日期', '2026/5/2'),
    ('建设单位', '常熟市虞创科技产业发展有限公司'),
    ('监理单位', '常熟市诚建工程监理有限公司'),
    ('设计单位', '中国建筑西南设计研究院有限公司（方案设计）、苏州安省建筑设计有限公司（施工设计）'),
    ('合同金额', '29431.54万元'),
]

for label, value in basic_info:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：{value}')
    set_run_font(run, '宋体', 12)

doc.add_paragraph()

# 二、概算控制情况
p = doc.add_paragraph()
run = p.add_run('二、概算控制情况')
set_run_font(run, '黑体', 14, True)

p = doc.add_paragraph()
run = p.add_run('本项目严格执行概算控制，各项费用总体可控。项目一季度累计完成产值约13318万元，实际成本控制在目标成本范围内，整体节约目标成本约583万元，体现较好的成本管控能力。')
set_run_font(run, '宋体', 12)

doc.add_paragraph()

# 三、成本分析汇总
p = doc.add_paragraph()
run = p.add_run('三、成本分析汇总')
set_run_font(run, '黑体', 14, True)

doc.add_paragraph()

# 表格
headers = ['序号', '费用项目', '合同收入(万元)', '目标成本(万元)', '实际成本(万元)', '目标成本节超(万元)', '盈亏(万元)', '目标成本节超率', '盈利率']
data = [
    ['1', '人工费', '1543.64', '1834.81', '1570.54', '+264.27', '-26.90', '16.83%', '-1.74%'],
    ['2', '直接材料费', '2898.20', '2572.74', '2560.44', '+12.30', '+337.76', '0.48%', '11.65%'],
    ['3', '专业分包工程费', '6600.27', '6009.28', '5721.33', '+287.95', '+878.93', '5.03%', '13.32%'],
    ['4', '措施项目费用', '1106.98', '1075.79', '1008.13', '+67.66', '+98.85', '6.71%', '8.93%'],
    ['5', '总包服务费', '0', '0', '0', '0', '0', '0', '0'],
    ['6', '签证变更费用', '0', '0', '0', '0', '0', '0', '0'],
    ['7', '现场经费', '902.48', '663.47', '706.05', '-42.58', '+196.44', '-6.03%', '21.76%'],
    ['8', '规费及其他应缴费', '266.85', '1291.60', '1231.81', '+59.79', '-964.96', '4.85%', '-361.61%'],
    ['9', '资金占用费', '0', '0', '30.89', '-30.89', '-30.89', '-100%', '—'],
    ['10', '税金及附加', '0', '60.60', '95.86', '-35.26', '-95.86', '-58.19%', '—'],
    ['11', '自行部分合计（不含税）', '13318.42', '13508.28', '12925.04', '+583.24', '+393.38', '4.51%', '2.95%'],
]

table = doc.add_table(rows=len(data)+1, cols=len(headers))
table.style = 'Table Grid'

# 设置表头
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    set_run_font(run, '宋体', 10, True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, 'D9D9D9')
    set_cell_border(cell, top={'val': 'single', 'sz': '4', 'color': '000000'},
                   bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                   left={'val': 'single', 'sz': '4', 'color': '000000'},
                   right={'val': 'single', 'sz': '4', 'color': '000000'})

# 数据行
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx+1].cells[col_idx]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        set_run_font(run, '宋体', 9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, top={'val': 'single', 'sz': '4', 'color': '000000'},
                       bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                       left={'val': 'single', 'sz': '4', 'color': '000000'},
                       right={'val': 'single', 'sz': '4', 'color': '000000'})

doc.add_paragraph()

# 四、各项费用详细分析
p = doc.add_paragraph()
run = p.add_run('四、各项费用详细分析')
set_run_font(run, '黑体', 14, True)

sections_content = [
    ('4.1 人工费分析', '合同收入1543.64万元，目标成本1834.81万元，实际成本1570.54万元，目标成本节约264.27万元，节约率14.46%，但利润为负26.90万元。主要原因：项目前期地下室主体结构阶段成本较高。'),
    ('4.2 直接材料费分析', '合同收入2898.20万元，目标成本2572.74万元，实际成本2560.44万元，节约12.30万元，节约率0.48%，盈利337.76万元，盈利率11.65%。'),
    ('4.3 专业分包工程费分析', '合同收入6600.27万元，目标成本6009.28万元，实际成本5721.33万元，节约287.95万元，节约率5.03%，盈利878.93万元，盈利率13.32%。'),
    ('4.4 措施项目费用分析', '合同收入1106.98万元，目标成本1075.79万元，实际成本1008.13万元，节约67.66万元，节约率6.71%，盈利98.85万元，盈利率8.93%。其中安全文明施工费超支92.04万元。'),
    ('4.5 现场经费分析', '合同收入902.48万元，目标成本663.47万元，实际成本706.05万元，超支42.58万元，盈利196.44万元。'),
    ('4.6 规费及其他应缴费分析', '实际成本1231.81万元，较目标成本节约59.79万元，但因合同收入仅为266.85万元，盈亏-964.96万元，主要因规费计费基数较大。'),
]

for title, content in sections_content:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, '黑体', 12, True)
    p = doc.add_paragraph()
    run = p.add_run(content)
    set_run_font(run, '宋体', 12)

doc.add_paragraph()

# 五、成本策划创效情况
p = doc.add_paragraph()
run = p.add_run('五、成本策划创效情况')
set_run_font(run, '黑体', 14, True)

p = doc.add_paragraph()
run = p.add_run('项目共制定30项成本策划措施，预期创效5087.19万元，实际完成1168.27万元。主要创效项包括量差策划（681.23万元）、专业分包招采策划（1644.72万元）等。')
set_run_font(run, '宋体', 12)

doc.add_paragraph()

# 六、存在的问题及改进措施
p = doc.add_paragraph()
run = p.add_run('六、存在的问题及改进措施')
set_run_font(run, '黑体', 14, True)

problems = [
    ('问题一：', '安全文明施工费超支约92万元，需加强现场安全文明施工管理，合理控制相关费用支出。'),
    ('问题二：', '规费及其他应缴费因计费基数较大导致亏损较大，建议后续在投标阶段充分考虑规费计取。'),
    ('问题三：', '现场经费超支42.58万元，需进一步优化项目管理流程，提高管理效率。'),
]

for label, content in problems:
    p = doc.add_paragraph()
    run = p.add_run(label + content)
    set_run_font(run, '宋体', 12)

doc.add_paragraph()

# 七、下一步工作计划
p = doc.add_paragraph()
run = p.add_run('七、下一步工作计划')
set_run_font(run, '黑体', 14, True)

plans = [
    '持续加强成本管控，确保各项费用控制在目标成本范围内。',
    '加快剩余分包工程的招标工作，确保工期进度。',
    '进一步推进成本策划创效措施的落地执行，争取完成更多创效目标。',
    '加强安全文明施工管理，控制安全文明施工费支出。',
]

for plan in plans:
    p = doc.add_paragraph()
    run = p.add_run(f'• {plan}')
    set_run_font(run, '宋体', 12)

# 保存
output_path = r'C:\Users\changhaoyan\Desktop\新建精密自动化产业园项目一季度全面分析报告.docx'
doc.save(output_path)
print(f"Word文档创建成功：{output_path}")
